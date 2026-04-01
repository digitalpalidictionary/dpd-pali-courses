"""
Verification tool for DOCX content integrity.
Compares text extracted from generated Word documents with source Markdown.
"""
import os
import sys
import re
from docx import Document
import yaml
from tools.printer import printer as pr

def get_docx_text(path):
    """Extracts text from a .docx file."""
    doc = Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

def verify_docx(docx_path, original_md_paths):
    """Verifies that the docx contains content from all original markdown files."""
    if not os.path.exists(docx_path):
        pr.warning(f"{docx_path} not found.")
        return False
        
    docx_text = get_docx_text(docx_path)
    # Basic cleaning of docx text for comparison
    docx_text_clean = docx_text.replace("\n", " ").replace("  ", " ")
    # Strip footnote markers from docx text if they were not converted
    docx_text_clean = re.sub(r'\[\^[^\]]+\]', '', docx_text_clean)
    
    missing_files = []
    for md_path in original_md_paths:
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
            
            # Extract potential text fragments for verification
            # 1. Headers
            headers = [ln.strip("# ").strip() for ln in content.split("\n") if ln.strip().startswith("#")]
            # 2. Paragraphs
            paragraphs = [ln.strip() for ln in content.split("\n") if len(ln.strip()) > 15 and not ln.strip().startswith(("<", "|", "#"))]
            # 3. Table cells (for table-only files)
            cells = []
            if "|" in content:
                for line in content.split("\n"):
                    if "|" in line:
                        parts = [p.strip() for p in line.split("|") if p.strip() and not all(c in "-:| " for c in p.strip())]
                        cells.extend([p for p in parts if len(p) > 10])

            # Combine all candidates, prioritizing paragraphs then headers then cells
            candidates = paragraphs + headers + cells
            
            found = False
            # Check up to 15 candidates for robustness
            for cand in candidates[:15]:
                # Remove MD attributes like {: #id}
                clean_cand = re.sub(r'\{: #[^}]+\}', '', cand)
                # Remove MD formatting: links, bold, italics, footnotes
                clean_cand = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_cand) # links
                clean_cand = clean_cand.replace('**', '').replace('*', '').replace('__', '').replace('_', '')
                clean_cand = re.sub(r'\[\^[^\]]+\]', '', clean_cand) # footnotes
                clean_cand = re.sub(r'^\s*(\d+\.|[-*+])\s+', '', clean_cand) # list markers
                clean_cand = clean_cand.replace("<br>", "").replace("&nbsp;", " ").replace("<br/>", "").strip()
                
                # Check if this cleaned candidate is in the DOCX
                if len(clean_cand) > 8 and clean_cand in docx_text_clean:
                    found = True
                    break
            
            if not found and candidates:
                missing_files.append(md_path)
                
    if missing_files:
        for f in missing_files:
            pr.warning(f"{os.path.basename(docx_path)}: missing content from {f}")
        return False
    return True

def main():
    docs_dir = "docs"
    docx_dir = "docx_exports"
    
    # We'll use the same logic as the generator to find files
    mkdocs_yaml_path = "mkdocs.yaml"
    with open(mkdocs_yaml_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)
        
    def extract_nav(item, base_dir, file_list):
        if isinstance(item, str):
            if item.endswith('.md'):
                file_list.append(os.path.join(base_dir, item))
        elif isinstance(item, list):
            for i in item:
                extract_nav(i, base_dir, file_list)
        elif isinstance(item, dict):
            for v in item.values():
                extract_nav(v, base_dir, file_list)

    all_files = []
    extract_nav(config.get("nav", []), docs_dir, all_files)
    
    folders = ['bpc', 'bpc_ex', 'bpc_key', 'ipc', 'ipc_ex', 'ipc_key']
    f_by_dir = {f: [] for f in folders}
    
    for file_path in all_files:
        rel = os.path.relpath(file_path, docs_dir)
        parts = rel.split(os.sep)
        folder = parts[0]
        if folder in f_by_dir and os.path.exists(file_path):
            # Consistent with generate_docx.py: skip folder-level index.md for _ex and _key
            if folder.endswith(('_ex', '_key')) and len(parts) == 2 and parts[1] == 'index.md':
                continue
            f_by_dir[folder].append(file_path)
            
    pr.green("Verifying DOCX content")
    all_success = True
    for folder, files in f_by_dir.items():
        docx_path = os.path.join(docx_dir, f"{folder}.docx")
        if os.path.exists(docx_path):
            if not verify_docx(docx_path, files):
                all_success = False

    if all_success:
        pr.yes("ok")
    else:
        pr.no("failures found")
        sys.exit(1)

if __name__ == '__main__':
    main()
