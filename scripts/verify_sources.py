"""
Interactive source verification script for DPD Pāḷi Courses.
Compares original DOCX materials against generated DOCX and PDF outputs.
"""

import os
import re
import json
import unicodedata
import argparse
import sys
import subprocess
from typing import List, Dict, Set

# Platform-safe keystroke detection
try:
    import tty as _tty_probe
    import termios as _termios_probe
    del _tty_probe, _termios_probe
    _HAS_TERMIOS = True
except ImportError:
    _HAS_TERMIOS = False
    try:
        import msvcrt as _msvcrt_probe
        del _msvcrt_probe
        _HAS_MSVCRT = True
    except ImportError:
        _HAS_MSVCRT = False

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pdfminer.high_level import extract_text as extract_pdf_text
from rich.console import Console
from rich.panel import Panel

console = Console()

def normalize_text(text: str) -> str:
    """Normalize text for robust comparison (stripping diacritics, sandhi markers, etc)."""
    if not text:
        return ""
    # Normalize to NFD to separate diacritics
    text = unicodedata.normalize('NFD', text)
    # Remove combining marks (diacritics)
    text = "".join([c for c in text if not unicodedata.combining(c)])
    text = text.lower()
    # Remove sandhi markers and common punctuation
    text = re.sub(r"['’‘]", '', text)
    text = re.sub(r'[^a-z0-9]', ' ', text)
    # Collapse whitespace
    text = " ".join(text.split())
    return text

def extract_docx_phrases(docx_path: str) -> List[Dict]:
    """Extract phrases from a DOCX file with context (heading and page estimate) in exact document order."""
    data = []
    try:
        doc = docx.Document(docx_path)
        current_heading = "Start of Document"
        
        # Simple heuristic for page numbers
        para_count = 0
        
        # Iterate over all elements in the body in order
        for block in doc.element.body:
            if isinstance(block, CT_P):
                para = Paragraph(block, doc)
                para_count += 1
                text = para.text.strip()
                
                if not text:
                    continue
                    
                if para.style is not None and para.style.name is not None and para.style.name.startswith('Heading'):
                    current_heading = text
                    
                norm = normalize_text(text)
                if len(norm) > 4:
                    data.append({
                        "orig": text,
                        "norm": norm,
                        "heading": current_heading,
                        "page": (para_count // 30) + 1
                    })
                    
            elif isinstance(block, CT_Tbl):
                table = Table(block, doc)
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if not text:
                            continue
                        norm = normalize_text(text)
                        if len(norm) > 4:
                            data.append({
                                "orig": text,
                                "norm": norm,
                                "heading": current_heading,
                                "page": (para_count // 30) + 1
                            })
                            
    except Exception as e:
        console.print(f"[red]Error reading DOCX {docx_path}: {e}[/red]")
    
    return data

def extract_docx_full_text(docx_path: str) -> str:
    """Extract all text from a DOCX for existence checking."""
    full_text = []
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
    except Exception as e:
        console.print(f"[red]Error reading target DOCX {docx_path}: {e}[/red]")
    return " ".join(full_text)

def extract_pdf_text_robust(pdf_path: str) -> str:
    """Extract text from PDF using pdftotext (if available) or pdfminer."""
    try:
        result = subprocess.run(['pdftotext', pdf_path, '-'], capture_output=True, text=True, check=True)
        return result.stdout
    except (subprocess.CalledProcessError, FileNotFoundError):
        try:
            return extract_pdf_text(pdf_path)
        except Exception as e:
            console.print(f"[red]Error extracting PDF {pdf_path}: {e}[/red]")
            return ""

def is_fuzzy_match(norm_phrase: str, target_text: str, threshold: float = 0.6) -> bool:
    """Check if a significant portion of words in the phrase exist in the target text."""
    words = [w for w in norm_phrase.split() if len(w) > 3]
    if not words:
        return False
    
    found_count = 0
    for word in words:
        if word in target_text:
            found_count += 1
            
    return (found_count / len(words)) >= threshold

def load_ignored(path: str) -> Dict[str, Set[str]]:
    """Load the skip list from JSON (returns a dict of heading -> set of normalized phrases)."""
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, list):
                    return {"Global": set(data)}
                elif isinstance(data, dict):
                    if any(k.endswith('.docx') for k in data.keys()):
                        all_phrases = set()
                        for phrases in data.values():
                            all_phrases.update(phrases)
                        return {"Global": all_phrases}
                    result = {}
                    for heading, phrases in data.items():
                        result[normalize_text(heading)] = set(phrases)
                    return result
        except Exception as e:
            console.print(f"[yellow]Warning: Could not load {path}: {e}[/yellow]")
    return {}

def save_ignored(path: str, ignored_data: Dict[str, Set[str]]):
    """Save the skip list to JSON."""
    try:
        dirpath = os.path.dirname(path)
        if dirpath:
            os.makedirs(dirpath, exist_ok=True)
        serializable_data = {k: sorted(list(v)) for k, v in ignored_data.items()}
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(serializable_data, f, indent=2, ensure_ascii=False)
    except Exception as e:
        console.print(f"[red]Error saving {path}: {e}[/red]")

def get_keystroke():
    """Wait for a single keystroke and return a semantic action."""
    if _HAS_TERMIOS:
        import tty
        import termios
        fd = sys.stdin.fileno()
        old_settings = termios.tcgetattr(fd)
        try:
            tty.setraw(sys.stdin.fileno())
            ch = sys.stdin.read(1)
        finally:
            termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
    elif _HAS_MSVCRT:
        import msvcrt
        ch = msvcrt.getch().decode('utf-8', errors='ignore')  # type: ignore[attr-defined]  # Windows-only API absent from macOS stubs
    else:
        # Fallback to standard input if no raw mode available
        ch = sys.stdin.read(1)
    
    # Map keys to actions
    if ch in ['\r', '\n', ' ']: # Enter or Space
        return "skip"
    elif ch == '1':
        return "ok"
    elif ch.lower() == 'q':
        return "quit"
    return None

def verify_source(source_path: str, docx_target: str, pdf_target: str, ignored_data: Dict[str, Set[str]], ignored_file: str, check_pdf: bool = True):
    """Verify a single source against its targets."""
    source_name = os.path.basename(source_path)
    console.print(Panel(f"[bold blue]Verifying Source:[/bold blue] {source_name}"))
    
    phrases_data = extract_docx_phrases(source_path)
    if not phrases_data:
        console.print("[yellow]No phrases extracted from source.[/yellow]")
        return ignored_data, False

    # Extract target texts
    target_docx_text = normalize_text(extract_docx_full_text(docx_target))
    target_docx_text_no_space = target_docx_text.replace(' ', '')
    
    pdf_text = ""
    pdf_text_no_space = ""
    if check_pdf and os.path.exists(pdf_target):
        pdf_text = normalize_text(extract_pdf_text_robust(pdf_target))
        pdf_text_no_space = pdf_text.replace(' ', '')
    
    new_ignored = {}
    for k, v in ignored_data.items():
        new_ignored[k] = set(v)
    current_session_ignored = set()

    for item in phrases_data:
        orig = item["orig"]
        norm = item["norm"]
        heading = item["heading"]
        page = item["page"]
        
        # Check if already ignored
        global_ignores = new_ignored.get("Global", set())
        heading_ignores = new_ignored.get(normalize_text(heading), set())
        if norm in global_ignores or norm in heading_ignores or norm in current_session_ignored:
            continue
            
        # Check existence in DOCX
        norm_no_space = norm.replace(' ', '')
        missing_docx = norm not in target_docx_text and norm_no_space not in target_docx_text_no_space
        
        # Check existence in PDF
        missing_pdf = False
        if check_pdf and pdf_target:
            if norm not in pdf_text and norm_no_space not in pdf_text_no_space:
                if not is_fuzzy_match(norm, pdf_text):
                    missing_pdf = True

        # Auto-approve if missing in PDF but present in DOCX
        if missing_pdf and not missing_docx:
            continue

        if missing_docx:
            target_types = ["DOCX"]
            if missing_pdf:
                target_types.append("PDF")
            
            joined = ", ".join(target_types)
            console.print(f"\n[bold red]MISSING[/bold red] in {joined}:")
            console.print(f"[cyan]Context:[/cyan] {heading} (approx. page {page})")
            console.print(f"[dim]Phrase:[/dim] [italic]\"{orig}\"[/italic]")
            
            prompt_parts = ["[white](Enter)[/white] Skip & Remember", "[white](1)[/white] OK now", "[white](q)[/white] Quit"]
            prompt_str = ", ".join(prompt_parts)
            console.print(f"[bold cyan]Action:[/bold cyan] {prompt_str}", end="")
            
            while True:
                action = get_keystroke()
                if action in ["skip", "ok", "quit"]:
                    break
            
            if action == "skip":
                norm_heading = normalize_text(heading)
                if norm_heading not in new_ignored:
                    new_ignored[norm_heading] = set()
                new_ignored[norm_heading].add(norm)
                console.print("\n[green]Added to permanent skip list (bound to heading).[/green]")
                # Save immediately to prevent data loss on Ctrl+C
                save_ignored(ignored_file, new_ignored)
            elif action == "ok":
                current_session_ignored.add(norm)
                console.print("\n[yellow]Skipped for this session.[/yellow]")
            elif action == "quit":
                console.print("\n[bold yellow]Quitting verification...[/bold yellow]")
                return new_ignored, True

    return new_ignored, False

def main():
    parser = argparse.ArgumentParser(description="Interactive Source Verification")
    parser.add_argument("--sources_dir", default="sources")
    parser.add_argument("--docx_dir", default="docx_exports")
    parser.add_argument("--pdf_dir", default="pdf_exports")
    parser.add_argument("--ignored_file", default="tests/ignored_sources.json")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF verification")
    args = parser.parse_args()

    ignored_data = load_ignored(args.ignored_file)
    
    mapping = {
        'Beginner_Pāli_Course.docx': 'bpc',
        'BPC_Exercises.docx': 'bpc_ex',
        'BPC_Key_to_Exercises.docx': 'bpc_key',
        'Intermediate_Pāli_Course.docx': 'ipc',
        'IPC_Exercises.docx': 'ipc_ex',
        'IPC_Key_to_Exercises.docx': 'ipc_key',
    }
    
    source_files = [f for f in os.listdir(args.sources_dir) if f.endswith('.docx')]
    
    for source_file in sorted(source_files):
        if source_file not in mapping:
            continue

        export_base = mapping[source_file]
        source_path = os.path.join(args.sources_dir, source_file)
        docx_target = os.path.join(args.docx_dir, f"{export_base}.docx")
        pdf_target = os.path.join(args.pdf_dir, f"{export_base}.pdf")
        
        updated_ignored, wants_quit = verify_source(source_path, docx_target, pdf_target, ignored_data, args.ignored_file, not args.no_pdf)
        
        ignored_data = updated_ignored
        save_ignored(args.ignored_file, ignored_data)
        
        if wants_quit:
            break
        
        console.print()

    console.print("\n[bold green]Verification complete.[/bold green]")

if __name__ == "__main__":
    main()
